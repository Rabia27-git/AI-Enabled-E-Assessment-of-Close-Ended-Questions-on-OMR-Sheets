from ultralytics import YOLO
from PIL import Image
import numpy as np
import cv2
import os
import pandas as pd
from scipy.stats import pointbiserialr
import requests
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from docx.shared import Inches
import matplotlib.pyplot as plt

# ---------------- CNN IMPORTS ----------------
import torch
import torch.nn as nn
from torchvision import transforms

# ---------------- BARCODE READER ----------------
from PIL import Image, ImageFilter
import zxingcpp

def extract_seat_number_from_barcode(file_path):
    """
    Reads barcode from TIFF (supports multi-page)
    Returns seat number string or None
    """

    try:
        with Image.open(file_path) as img:

            for i in range(getattr(img, 'n_frames', 1)):
                img.seek(i)

                page = img.convert('L')

                # Upscale for better detection
                w, h = page.size
                page = page.resize((w*3, h*3), resample=Image.Resampling.LANCZOS)

                # Sharpen image
                page = page.filter(ImageFilter.SHARPEN)

                results = zxingcpp.read_barcodes(page)

                if results:
                    for res in results:
                        return res.text.strip()   

        return None 

    except Exception as e:
        print(f"Barcode error: {e}")
        return None

# ---------------- CONFIG ----------------
QB_MODEL_PATH = "runs/detect/question_block_train/weights/best.pt"
BUBBLE_MODEL_PATH = "runs/detect/bubble_train_finetune/weights/best.pt"
CNN_MODEL_PATH = "bubble_classifier.pth"
ANSWER_KEY_PATH = "answer key.xlsx"

TEST_FOLDER = ""
OUTPUT_FOLDER = "Output"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# ---------------- PSYCHOMETRIC REPORT GENERATION ----------------
def generate_psychometric_report():

    # ---------------- CONFIG ----------------
    OPENROUTER_API_KEY = ""

    RESPONSES_FILE = os.path.join(OUTPUT_FOLDER, "Result.xlsx")
    OUTPUT_DOC = os.path.join(OUTPUT_FOLDER, "Psychometric_Exam_Report.docx")
    ANSWER_KEY_FILE = ANSWER_KEY_PATH   # reuse existing

    MODEL_NAME = "openai/gpt-4o-mini"

    # ---------------- LOAD DATA ----------------

    responses = pd.read_excel(RESPONSES_FILE)
    answer_key = pd.read_excel(ANSWER_KEY_FILE)

    question_cols = [c for c in responses.columns if str(c).startswith("Q")]
    responses_subset = responses[question_cols]

    # ---------------- SCORE RESPONSES ----------------

    def score_responses(df, keys):

        scored = pd.DataFrame()

        for q in df.columns:

            q_num = int(q.replace("Q", ""))

            correct_val = keys.loc[
                keys["QuestionNo"] == q_num,
                "CorrectOption"
            ].values[0]

            correct_set = set(a.strip().upper() for a in str(correct_val).split(','))

            def is_correct(resp):

                if pd.isna(resp):
                    return 0

                resp_set = set(a.strip().upper() for a in str(resp).split(','))

                return 1 if resp_set == correct_set else 0

            scored[q] = df[q].apply(is_correct)

        return scored


    scored = score_responses(responses_subset, answer_key)

    # ---------------- BASIC TEST STATS ----------------

    total_students = len(responses)
    k = scored.shape[1]

    # ---------------- ITEM DIFFICULTY ----------------

    difficulty = scored.mean()

    diff_summary = {
        "Easy": int((difficulty > 0.70).sum()),
        "Medium": int(((difficulty >= 0.30) & (difficulty <= 0.70)).sum()),
        "Hard": int((difficulty < 0.30).sum())
    }

    # ---------------- ITEM DISCRIMINATION ----------------

    total_scores = scored.sum(axis=1)

    disc_results = []

    for q in scored.columns:

        corrected_total = total_scores - scored[q]

        if scored[q].nunique() <= 1:
            r = 0
        else:
            r, _ = pointbiserialr(scored[q], corrected_total)

        disc_results.append(round(float(r), 3))

    disc_series = pd.Series(disc_results)

    disc_summary = {
        "Excellent": int((disc_series >= 0.40).sum()),
        "Good": int(((disc_series >= 0.30) & (disc_series < 0.40)).sum()),
        "Fair": int(((disc_series >= 0.20) & (disc_series < 0.30)).sum()),
        "Poor": int((disc_series < 0.20).sum())
    }

    # ---------------- RELIABILITY (CRONBACH ALPHA) ----------------

    item_vars = scored.var(axis=0, ddof=1).sum()
    total_var = total_scores.var(ddof=1)

    if total_var == 0:
        alpha = 0
    else:
        alpha = (k / (k - 1)) * (1 - (item_vars / total_var))

    alpha = round(alpha, 3)

    # ---------------- DISTRACTOR ANALYSIS ----------------

    nfd_total = 0
    nfd_items = []  # Store NFD info for table

    for q in responses_subset.columns:
        q_num = int(q.replace("Q", ""))
        correct_val = str(
            answer_key.loc[
                answer_key["QuestionNo"] == q_num,
                "CorrectOption"
            ].values[0]
        ).strip().upper()

        counts = (
            responses_subset[q]
            .astype(str)
            .str.strip()
            .str.upper()
            .value_counts(normalize=True) * 100
        )

        nfd_opts = []
        # for opt, pct in counts.items():
        #     if opt != correct_val and pct < 5:
        #         nfd_total += 1
        #         nfd_opts.append(opt)
        for opt, pct in counts.items():
            if opt == "-":
                continue

            if opt != correct_val and pct < 5:
                nfd_total += 1
                nfd_opts.append(opt)

        nfd_items.append({
            "Question": q,
            "NonFunctionalDistractors": ', '.join(nfd_opts) if nfd_opts else "None"
        })

    # ---------------- PROMPT ----------------

    prompt = f"""
    You are a professional psychometrician writing a formal examination analysis report.

    Exam Statistics:

    Total Students: {total_students}
    Total Items: {k}

    Cronbach Alpha: {alpha}

    Item Difficulty Distribution:
    {json.dumps(diff_summary, indent=2)}

    Item Discrimination Distribution:
    {json.dumps(disc_summary, indent=2)}

    Number of Non Functional Distractors: {nfd_total}

    Write a formal psychometric item analysis report.

    Structure:

    1 Executive Summary
    2 Methodology
    3 Item Difficulty Analysis
    4 Item Discrimination Analysis
    5 Reliability Analysis
    6 Distractor Analysis
    7 Recommendations

    Use professional academic language suitable for research reports.

    Return STRICT JSON format:

    {{
    "executive_summary":"",
    "methodology":"",
    "difficulty_analysis":"",
    "discrimination_analysis":"",
    "reliability_analysis":"",
    "distractor_analysis":"",
    "recommendations":""
    }}

    Include a concluding remark that summarizes overall test quality, which will align with a visual summary figure presented at the end of the report.
    """

    # ---------------- CALL MODEL ----------------

    url = "https://openrouter.ai/api/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system",
            "content": "You are an expert psychometrician and educational assessment researcher."},

            {"role": "user",
            "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 1500
    }

    response = requests.post(url, headers=headers, json=data)

    result = response.json()

    # DEBUG: print response if something goes wrong
    if "choices" not in result:
        print("API Error Response:")
        print(result)
        raise Exception("Claude API request failed.")

    content = result["choices"][0]["message"]["content"]

    # ----- Extract JSON safely -----
    start = content.find("{")
    end = content.rfind("}") + 1

    if start == -1 or end == -1:
        print("Claude returned invalid response:")
        print(content)
        raise Exception("JSON extraction failed")

    json_str = content[start:end]

    report = json.loads(json_str)

    SUMMARY_IMAGE = os.path.join(OUTPUT_FOLDER, "summary.png")

    def generate_summary_image():
        
        fig, axs = plt.subplots(2, 2, figsize=(10, 8))

        # ---- Difficulty Pie ----
        axs[0, 0].pie(
            diff_summary.values(),
            labels=diff_summary.keys(),
            autopct='%1.1f%%'
        )
        axs[0, 0].set_title("Item Difficulty")

        # ---- Discrimination Pie ----
        axs[0, 1].pie(
            disc_summary.values(),
            labels=disc_summary.keys(),
            autopct='%1.1f%%'
        )
        axs[0, 1].set_title("Item Discrimination")

        # ---- Reliability Text ----
        axs[1, 0].axis('off')
        axs[1, 0].text(
            0.5, 0.5,
            f"Cronbach Alpha\n{alpha}",
            ha='center', va='center',
            fontsize=14
        )

        # ---- Distractor Info ----
        axs[1, 1].axis('off')
        axs[1, 1].text(
            0.5, 0.5,
            f"Non-Functional Distractors\n{nfd_total}",
            ha='center', va='center',
            fontsize=14
        )

        plt.tight_layout()
        plt.savefig(SUMMARY_IMAGE, dpi=300, bbox_inches='tight')
        plt.close()

    # CALL IT
    generate_summary_image()

    # ---------------- CREATE WORD DOCUMENT ----------------

    doc = Document()

    title = doc.add_heading("Post-Hoc Psychometric Examination Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(report["executive_summary"])

    # Methodology
    doc.add_heading("2. Methodology", level=1)
    doc.add_paragraph(report["methodology"])

    # Difficulty
    doc.add_heading("3. Item Difficulty Analysis", level=1)
    doc.add_paragraph(report["difficulty_analysis"])

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Range"
    hdr[2].text = "Items"
    hdr[3].text = "Interpretation"

    difficulty_rows = [
    ("Easy", ">0.70", diff_summary["Easy"], "Items answered correctly by most students."),
    ("Medium", "0.30-0.70", diff_summary["Medium"], "Moderate difficulty level."),
    ("Hard", "<0.30", diff_summary["Hard"], "Items answered correctly by few students.")
    ]

    for name, rng, count, interp in difficulty_rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = rng
        row[2].text = str(count)
        row[3].text = interp

    # Discrimination
    doc.add_heading("4. Item Discrimination Analysis", level=1)
    doc.add_paragraph(report["discrimination_analysis"])

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Range"
    hdr[2].text = "Items"

    disc_rows = [
    ("Excellent", "≥0.40", disc_summary["Excellent"]),
    ("Good", "0.30-0.39", disc_summary["Good"]),
    ("Fair", "0.20-0.29", disc_summary["Fair"]),
    ("Poor", "<0.20", disc_summary["Poor"])
    ]

    for name, rng, count in disc_rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = rng
        row[2].text = str(count)

    # Reliability
    doc.add_heading("5. Reliability Analysis", level=1)
    doc.add_paragraph(report["reliability_analysis"])

    doc.add_paragraph(f"Cronbach Alpha Value: {alpha}")

    # Distractor
    doc.add_heading("6. Distractor Analysis", level=1)
    doc.add_paragraph(report["distractor_analysis"])
    doc.add_paragraph(f"Total Non Functional Distractors Detected: {nfd_total}")

    nfd_items_filtered = [item for item in nfd_items if item["NonFunctionalDistractors"] not in ["None", "-"]]

    # Add NFD Table (only with real distractors)
    doc.add_paragraph("Table: Non-Functional Distractors by Item")
    if nfd_items_filtered:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Non-Functional Distractors"

        for item in nfd_items_filtered:
            row = table.add_row().cells
            row[0].text = item["Question"]
            row[1].text = item["NonFunctionalDistractors"]
    else:
        doc.add_paragraph("No non-functional distractors detected in this test.")

    # Recommendations
    doc.add_heading("7. Recommendations", level=1)
    doc.add_paragraph(report["recommendations"])

    doc.add_page_break()
    doc.add_heading("8. Visual Summary of Psychometric Analysis", level=1)
    doc.add_paragraph(
        "The following figure provides a consolidated visual overview of key psychometric indicators including item difficulty, item discrimination, test reliability, and distractor efficiency."
    )
    doc.add_picture(SUMMARY_IMAGE, width=Inches(6))

    # ---------------- FORMAT ----------------

    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    doc.save(OUTPUT_DOC)

    print("Report Generated Successfully:", OUTPUT_DOC)

# ---------------- CNN MODEL DEFINITION ----------------
class BubbleCNN(nn.Module):
    def __init__(self, num_classes=2):
        super(BubbleCNN, self).__init__()
        self.features = nn.Sequential(
            nn.Conv2d(3,32,3,padding=1),
            nn.ReLU(),
            nn.MaxPool2d(2),

            nn.Conv2d(32,64,3,padding=1),
            nn.ReLU(),
            nn.MaxPool2d(2),

            nn.Conv2d(64,128,3,padding=1),
            nn.ReLU(),
            nn.MaxPool2d(2),
        )

        self.classifier = nn.Sequential(
            nn.Flatten(),
            nn.Linear(128*4*4,128),
            nn.ReLU(),
            nn.Dropout(0.3),
            nn.Linear(128,2)
        )

    def forward(self,x):
        x=self.features(x)
        x=self.classifier(x)
        return x


# ---------------- LOAD CNN ----------------
cnn_model = BubbleCNN().to(DEVICE)
cnn_model.load_state_dict(torch.load(CNN_MODEL_PATH, map_location=DEVICE))
cnn_model.eval()

cnn_transform = transforms.Compose([
    transforms.ToPILImage(),
    transforms.Resize((32,32)),
    transforms.ToTensor(),
    transforms.Normalize([0.5,0.5,0.5],[0.5,0.5,0.5])
])

class_names = ["filled","unfilled"]


# ---------------- LOAD YOLO MODELS ----------------
qb_model = YOLO(QB_MODEL_PATH)
bubble_model = YOLO(BUBBLE_MODEL_PATH)


# ---------------- SIMPLE NMS FUNCTION ----------------
def nms(boxes, iou_thresh=0.3):
    if not boxes:
        return []
    boxes_array = np.array(boxes)
    x1 = boxes_array[:,0]
    y1 = boxes_array[:,1]
    x2 = boxes_array[:,2]
    y2 = boxes_array[:,3]
    areas = (x2 - x1) * (y2 - y1)
    order = np.argsort(x1)
    keep = []
    while order.size > 0:
        i = order[0]
        keep.append(i)
        xx1 = np.maximum(x1[i], x1[order[1:]])
        yy1 = np.maximum(y1[i], y1[order[1:]])
        xx2 = np.minimum(x2[i], x2[order[1:]])
        yy2 = np.minimum(y2[i], y2[order[1:]])
        w = np.maximum(0.0, xx2 - xx1)
        h = np.maximum(0.0, yy2 - yy1)
        inter = w * h
        iou = inter / (areas[i] + areas[order[1:]] - inter)
        inds = np.where(iou <= iou_thresh)[0]
        order = order[inds + 1]
    return [boxes[k] for k in keep]

def compute_fill_ratio(bubble_img):

    gray = cv2.cvtColor(bubble_img, cv2.COLOR_RGB2GRAY)

    # binarize
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

    dark_pixels = np.sum(thresh == 255)
    total_pixels = thresh.size

    fill_ratio = dark_pixels / total_pixels

    return fill_ratio


# Removing Duplicates (UNCHANGED)
def remove_duplicates_and_contained(boxes, iou_thresh=0.3, contain_thresh=0.8):

    if not boxes:
        return []

    boxes = np.array(boxes)
    x1 = boxes[:, 0]
    y1 = boxes[:, 1]
    x2 = boxes[:, 2]
    y2 = boxes[:, 3]

    areas = (x2 - x1) * (y2 - y1)
    order = np.argsort(areas)[::-1]

    keep = []
    removed = set()

    for i in order:
        if i in removed:
            continue

        keep.append(i)

        for j in order:
            if j == i or j in removed:
                continue

            xx1 = max(x1[i], x1[j])
            yy1 = max(y1[i], y1[j])
            xx2 = min(x2[i], x2[j])
            yy2 = min(y2[i], y2[j])

            w = max(0, xx2 - xx1)
            h = max(0, yy2 - yy1)
            inter = w * h

            if inter == 0:
                continue

            union = areas[i] + areas[j] - inter
            iou = inter / union if union > 0 else 0

            containment = inter / areas[j] if areas[j] > 0 else 0
            area_ratio = areas[j] / areas[i]

            if iou > iou_thresh or containment > contain_thresh or area_ratio < 0.6:
                removed.add(j)

    return [tuple(boxes[k]) for k in keep if k not in removed]


# ---------------- ANSWER KEY ----------------
answer_df = pd.read_excel(ANSWER_KEY_PATH)
answer_dict = {}
for idx, row in answer_df.iterrows():
    q_no = int(row['QuestionNo'])
    correct_opts = row['CorrectOption'].split(',')  # handle multiple correct options
    answer_dict[q_no] = [opt.strip() for opt in correct_opts]

# ---------------- EXCEL RESULT STORAGE ----------------
results_list = []

# ---------------- PROCESS TIFFs SAFELY ----------------
for file in os.listdir(TEST_FOLDER):

    if not file.lower().endswith((".tif",".tiff")):
        continue

    tiff_path = os.path.join(TEST_FOLDER, file)

    # ---------------- GET SEAT NUMBER FROM BARCODE ----------------
    seat_number = extract_seat_number_from_barcode(tiff_path)

    if seat_number is None:
        print(f"⚠ No barcode found for {file}, assigning UNKNOWN")
        seat_number = "UNKNOWN"

    q_no = 1
    sheet_answers = {}
    total_correct = 0
    total_incorrect = 0
    total_blank = 0
    tiff_path = os.path.join(TEST_FOLDER,file)

    try:

        pil_img = Image.open(tiff_path).convert("RGB")
        img = np.array(pil_img)

        qb_results = qb_model.predict(
            source=img,
            imgsz=1280,
            conf=0.40,
            save=False,
            verbose=False
        )

        boxes = []

        for result in qb_results:

            if result.boxes is None:
                continue

            for box in result.boxes:

                conf=float(box.conf[0])

                if conf<0.40:
                    continue

                x1,y1,x2,y2=map(int,box.xyxy[0])
                boxes.append((x1,y1,x2,y2,conf))

        if not boxes:
            print("No question blocks found")
            continue


        boxes_sorted_x=sorted(boxes,key=lambda b:b[0])

        columns=[]
        col_thresh=50

        for box in boxes_sorted_x:

            x1,_,_,_,_=box
            placed=False

            for col in columns:

                if abs(np.mean([b[0] for b in col])-x1)<col_thresh:

                    col.append(box)
                    placed=True
                    break

            if not placed:
                columns.append([box])

        for col in columns:
            col.sort(key=lambda b:b[1])


        total_bubbles_sheet=0

        total_questions_detected=len(boxes)
        expected_options=5 if total_questions_detected==100 else 4

        if total_questions_detected == 100:
            DARKNESS_THRESHOLD = 0.40
        elif total_questions_detected == 180:
            DARKNESS_THRESHOLD = 0.40
        elif total_questions_detected == 200:
            DARKNESS_THRESHOLD = 0.50
        else:
            DARKNESS_THRESHOLD = 0.45

        option_letters = ["A","B","C","D","E"]
        print(f"Detected {total_questions_detected} questions")


        for col in columns:

            for box in col:

                x1,y1,x2,y2,conf=box

                #cv2.rectangle(img,(x1,y1),(x2,y2),(0,255,0),3)

                crop=img[y1:y2,x1:x2]

                bubble_results=bubble_model.predict(
                    source=crop,
                    imgsz=640,
                    conf=0.20,
                    save=False,
                    verbose=False
                )

                detected_bubbles=[]

                for result in bubble_results:

                    if result.boxes is None:
                        continue

                    for bubble_box in result.boxes:

                        bx1,by1,bx2,by2=map(int,bubble_box.xyxy[0])
                        detected_bubbles.append((bx1,by1,bx2,by2))


                detected_bubbles=remove_duplicates_and_contained(detected_bubbles)

                if len(detected_bubbles)>expected_options:
                    detected_bubbles=sorted(detected_bubbles,key=lambda b:b[0])[-expected_options:]


                # SORT LEFT→RIGHT
                detected_bubbles=sorted(detected_bubbles,key=lambda b:b[0])

                total_bubbles_sheet += len(detected_bubbles)

                option_index = 0
                bubble_scores = {}

                # CNN stage storage
                cnn_candidates = []
                bubble_images = {}

                for bx1,by1,bx2,by2 in detected_bubbles:

                    # DRAW BLUE BOX (visual only)
                    cv2.rectangle(
                        img,
                        (x1+bx1,y1+by1),
                        (x1+bx2,y1+by2),
                        (255,0,0),
                        2
                    )

                    # ---------------- CNN CROP ----------------
                    bubble_crop = crop[by1:by2, bx1:bx2].copy()
                    bubble_crop = cv2.resize(bubble_crop, (32,32))

                    tensor = cnn_transform(bubble_crop).unsqueeze(0).to(DEVICE)

                    with torch.no_grad():
                        output = cnn_model(tensor)
                        probs = torch.softmax(output, dim=1)

                        filled_conf = probs[0][0].item()   # CNN probability of "filled"

                     # ---------------- DARKNESS SCORE ----------------
                    fill_ratio = compute_fill_ratio(bubble_crop)

                    if option_index < len(option_letters):
                        opt = option_letters[option_index]

                        bubble_images[opt] = bubble_crop

                        # CNN decides if candidate
                        if filled_conf > 0.5:
                            cnn_candidates.append(opt)

                    print(f"Q{q_no} bubble → CNN:{filled_conf:.2f} Fill:{fill_ratio:.2f}")

                    option_index += 1

                filled_bubbles = []
                if len(cnn_candidates) == 1:
                    filled_bubbles = cnn_candidates

                elif len(cnn_candidates) > 1:

                    for opt in cnn_candidates:

                        bubble_img = bubble_images[opt]
                        fill_ratio = compute_fill_ratio(bubble_img)

                        if fill_ratio >= DARKNESS_THRESHOLD:
                            filled_bubbles.append(opt)

                # ---------------- FINAL SELECTED OPTION ----------------
                correct_options = answer_dict.get(q_no, [])

                # NO bubble filled
                if len(filled_bubbles) == 0:

                    total_blank += 1
                    display_option = "-"
                    block_color = (0,0,255)
                    text_color = (0,0,255)

                # ONE bubble filled
                elif len(filled_bubbles) == 1:

                    filled_option = filled_bubbles[0]

                    if filled_option in correct_options:
                        total_correct += 1
                        block_color = (0,255,0)
                        text_color = (0,255,0)
                    else:
                        total_incorrect += 1
                        block_color = (0,0,255)
                        text_color = (0,0,255)

                    display_option = filled_option

                # MULTIPLE bubbles filled
                else:

                    filled_set = set(filled_bubbles)
                    correct_set = set(correct_options)

                    # correct if student filled correct answers only
                    if filled_set.issubset(correct_set):

                        total_correct += 1
                        block_color = (0,255,0)
                        text_color = (0,255,0)

                    else:

                        total_incorrect += 1
                        block_color = (0,0,255)
                        text_color = (0,0,255)

                    display_option = ",".join(filled_bubbles)

                sheet_answers[f"Q{q_no}"] = display_option

                # ---------------- DRAW RECTANGLE & TEXT ----------------
                cv2.rectangle(img, (x1,y1), (x2,y2), block_color, 3)
                text = f"Q{q_no}:{display_option}"
                cv2.putText(
                    img,
                    text,
                    (x1 - 60, y1 + 30), 
                    cv2.FONT_HERSHEY_SIMPLEX,
                    0.5,
                    text_color,
                    2
                )

                q_no+=1


        out_name=os.path.splitext(file)[0]+"_annotated.png"

        # Resize annotated image before saving
        scale_percent = 50   # reduce to 50%
        width = int(img.shape[1] * scale_percent / 100)
        height = int(img.shape[0] * scale_percent / 100)

        resized_img = cv2.resize(img, (width, height), interpolation=cv2.INTER_AREA)

        cv2.imwrite(os.path.join(OUTPUT_FOLDER,out_name), resized_img)

        # ---------------- SAVE RESULTS FOR THIS SHEET ----------------
        sheet_answers["SeatNumber"] = seat_number
        sheet_answers["Total_Correct"] = total_correct
        sheet_answers["Total_Incorrect"] = total_incorrect
        sheet_answers["Total_Blank"] = total_blank
        results_list.append(sheet_answers)

        print("Processed:",file)
        print("Questions:",q_no-1)
        print("Bubbles:",total_bubbles_sheet)
        print("------------")


    except Exception as e:
        print("Failed:",file,e)

# ---------------- AFTER ALL SHEETS ----------------
results_df = pd.DataFrame(results_list)

# Reorder columns: SeatNumber → Q1,Q2... → totals
all_q_cols = [f"Q{i}" for i in range(1, q_no)]
final_cols = ["SeatNumber"] + all_q_cols + ["Total_Correct","Total_Incorrect","Total_Blank"]
results_df = results_df.reindex(columns=final_cols)

results_df.to_excel(os.path.join(OUTPUT_FOLDER,"Result.xlsx"), index=False)
print("Excel result saved!")

print("DONE")
generate_psychometric_report()